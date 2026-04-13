#!/usr/bin/env python3
"""Generate Threat Model Word document matching the reference format."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0, 0, 0)

# Paragraph spacing
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

# --- Title ---
title = doc.add_heading('Threat Model: Guidance for AI-Powered Restaurant Visibility on AWS', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- Section 1 ---
doc.add_heading('1. What are we building?', level=1)

doc.add_paragraph(
    'Guidance for AI-Powered Restaurant Visibility on AWS is an AWS Guidance that demonstrates '
    'how to build an AI-powered restaurant visibility system using Amazon Bedrock '
    'AgentCore. It monitors kitchen equipment temperatures, inventory levels, and staffing '
    'across 10 Georgia restaurant locations, with text and voice chat interfaces for operators.'
)

doc.add_heading('Technical Components', level=2)

components = [
    ("Amazon CloudFront:", "CDN for frontend hosting with TLS 1.2 enforcement, security response headers (HSTS, X-Frame-Options: DENY), and Origin Access Control for S3. Deployed in AWS, managed by AWS."),
    ("AWS WAF:", "Web application firewall on CloudFront with AWS Managed Rules (Common Rule Set, Known Bad Inputs Rule Set). Deployed in AWS, managed by AWS."),
    ("Amazon S3:", "Static website hosting for the monitoring dashboard (HTML/JS/CSS). Block public access enabled, versioning, AES-256 encryption, access logging to dedicated logs bucket with 90-day lifecycle."),
    ("Amazon Cognito User Pool:", "User authentication with optional MFA (software tokens), admin-only registration, 12+ character password policy with complexity requirements. OAuth2 Authorization Code flow."),
    ("Amazon API Gateway:", "REST API with Cognito authorizer on all protected endpoints. Six data endpoints (/restaurants, /equipment, /inventory, /staffing, /tickets, /voice-token) plus /chat. Request validation enabled."),
    ("AWS Lambda (API Function):", "Python 3.13 serverless compute for data queries and voice token generation. 256MB memory, 30s timeout, reserved concurrency of 10. KMS-encrypted environment variables. SQS dead letter queue."),
    ("AWS Lambda (Chat Function):", "Python 3.13 serverless compute for AgentCore invocation. 512MB memory, 300s timeout, reserved concurrency of 10. KMS-encrypted environment variables. SQS dead letter queue."),
    ("Amazon Bedrock AgentCore Runtime:", "Executes the Strands SDK agent with Amazon Nova Lite model (us.amazon.nova-lite-v1:0). Nine tools for restaurant operations. Deployed in AWS, managed by AWS."),
    ("Amazon Bedrock AgentCore Memory:", "Short-term and long-term conversation memory (STM + LTM) for personalized operator interactions across sessions. Deployed in AWS, managed by AWS."),
    ("Amazon Nova Sonic:", "Bidirectional voice streaming via AgentCore WebSocket endpoint at 16kHz audio sample rate for hands-free kitchen interaction. Managed by AWS."),
    ("Amazon Bedrock Knowledge Base:", "RAG over 6 equipment manuals using Amazon Titan Embed Text v2 for vector embeddings. Connected to OpenSearch Serverless for vector search."),
    ("Amazon OpenSearch Serverless:", "Vector search backend for the Knowledge Base. VECTORSEARCH collection type with encryption and network security policies."),
    ("Amazon DynamoDB (6 tables):", "NoSQL storage for restaurants, equipment-readings, inventory-items, staffing-requirements, tickets, and chat-history. PAY_PER_REQUEST billing, KMS CMK encryption, Point-in-Time Recovery enabled on all tables. TTL on chat-history table."),
    ("AWS KMS (2 Customer Managed Keys):", "DynamoDB table encryption key and Lambda environment variable encryption key. Both with automatic annual rotation."),
    ("Amazon SQS:", "Dead letter queues for both Lambda functions. SSL-only access policies enforced."),
    ("Amazon CloudWatch:", "Logs and metrics for all Lambda functions, API Gateway, and DynamoDB tables."),
    ("AWS CloudTrail:", "API call auditing across all services."),
    ("AWS CloudFormation:", "Infrastructure as Code. Stacks: base infrastructure, knowledge base, chat endpoint."),
]

for bold_text, normal_text in components:
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run(bold_text + " ")
    run_bold.bold = True
    run_bold.font.size = Pt(11)
    run_normal = p.add_run(normal_text)
    run_normal.font.size = Pt(11)

doc.add_heading('Architecture', level=2)

doc.add_paragraph('The solution provides an AI-powered restaurant monitoring system with text and voice interfaces:')

arch_items = [
    ("Monitoring Dashboard:", "Real-time equipment temperatures, inventory levels, staffing schedules, and maintenance tickets across 10 locations"),
    ("AI Chat Agent:", "Strands-based agent that queries operational data, analyzes temperature anomalies, creates maintenance tickets, and searches equipment manuals"),
    ("Voice Interface:", "Bidirectional voice streaming via Nova Sonic for hands-free kitchen operation"),
]

for bold_text, normal_text in arch_items:
    p = doc.add_paragraph(style='List Bullet')
    run_bold = p.add_run(bold_text + " ")
    run_bold.bold = True
    run_normal = p.add_run(normal_text)

doc.add_paragraph('Data flows:')

flows = [
    "User → CloudFront (WAF) → S3 (static dashboard assets)",
    "User → CloudFront → API Gateway (Cognito authorizer) → API Lambda → DynamoDB (data queries)",
    "User → API Gateway (Cognito authorizer) → Chat Lambda → AgentCore Runtime → Strands Agent → DynamoDB (tool calls)",
    "Strands Agent → Bedrock Knowledge Base → OpenSearch Serverless (equipment manual RAG)",
    "User → AgentCore WebSocket → Nova Sonic (bidirectional voice)",
    "Equipment simulator → DynamoDB (temperature readings via load-data.sh)",
]

for i, flow in enumerate(flows, 1):
    doc.add_paragraph(f"{i}. {flow}", style='List Number')

doc.add_paragraph(
    'Deployment: CloudFormation templates deployed via deploy-all.sh shell script. '
    'AgentCore agent deployed separately via AgentCore CLI (agentcore deploy).'
)

# --- Section 2 ---
doc.add_heading('2. What can go wrong?', level=1)

threats = [
    {
        "title": "Unauthorized API Access",
        "risk": "Attackers could access the restaurant data API or chat endpoint without proper authentication",
        "impact": "Unauthorized access to equipment temperatures, inventory levels, staffing schedules, and ability to create maintenance tickets. Cost escalation from Bedrock model invocations",
        "likelihood": "Low — All data and chat endpoints are protected by Cognito User Pool authorizer. Admin-only user creation prevents self-registration",
    },
    {
        "title": "Stolen Cognito Token Replay",
        "risk": "Stolen or leaked JWT tokens could be replayed to access API endpoints",
        "impact": "Unauthorized data access across all 10 restaurant locations for the token's validity period",
        "likelihood": "Medium — Tokens are transmitted over HTTPS but could be extracted from browser storage or intercepted via XSS if CORS is misconfigured",
    },
    {
        "title": "Prompt Injection",
        "risk": "Malicious users could craft chat inputs to manipulate the Strands agent behavior, bypass tool routing, or extract system prompts",
        "impact": "Agent could create false maintenance tickets, return fabricated data, or disclose system configuration details",
        "likelihood": "Medium — The agent processes user input directly. System prompt defines boundaries but no Bedrock Guardrails are configured (no content filters, prompt attack detection, or PII filters)",
    },
    {
        "title": "Sensitive Data Exposure in Agent Responses",
        "risk": "Agent could inadvertently expose staffing PII (manager names), internal table names, or system configuration in chat responses",
        "impact": "Privacy violation, information disclosure enabling further attacks",
        "likelihood": "Low — Agent tools return only DynamoDB data. System prompt defines strict boundaries. Lambda error handler returns generic messages",
    },
    {
        "title": "CORS Wildcard Allows Cross-Origin Data Access",
        "risk": "Access-Control-Allow-Origin: * on all API responses allows any website to make authenticated requests to the API",
        "impact": "Cross-site data exfiltration if combined with token theft. Attacker-controlled website could read restaurant operational data",
        "likelihood": "Medium — CORS wildcard is present in both Lambda functions and API Gateway gateway responses",
    },
    {
        "title": "Demo Credentials in Deployment Script",
        "risk": "The deploy-all.sh script contains hardcoded demo user credentials (email and password) created via admin-set-user-password",
        "impact": "Anyone with access to the repository can log into the deployed system",
        "likelihood": "Medium — Credentials are visible in the deployment script. Only mitigated by the fact that the Cognito User Pool ID is deployment-specific",
    },
    {
        "title": "Denial of Service",
        "risk": "Excessive API requests could exhaust Lambda reserved concurrency (10), DynamoDB on-demand capacity, or Bedrock model quotas",
        "impact": "Dashboard unavailable for all 10 restaurant locations. Chat and voice interfaces unresponsive",
        "likelihood": "Medium — Lambda reserved concurrency acts as a ceiling but also limits legitimate traffic. No API Gateway throttling configured beyond WAF managed rule sets",
    },
    {
        "title": "IAM Wildcard on AgentCore Permissions",
        "risk": "Both Lambda roles use Resource: * for bedrock-agentcore:InvokeAgentRuntime permissions",
        "impact": "Lambda functions could invoke any AgentCore runtime in the account, not just the restaurant agent",
        "likelihood": "Low — Exploitation requires compromising the Lambda execution environment. Agent ARN is dynamic at deploy time",
    },
    {
        "title": "Knowledge Base Data Tampering",
        "risk": "Tampered equipment manuals in the KB S3 bucket could provide incorrect troubleshooting guidance",
        "impact": "Operators following incorrect procedures could damage equipment or create safety hazards",
        "likelihood": "Low — S3 bucket has public access blocked, SSL-only policy, and versioning enabled",
    },
    {
        "title": "Conversation Memory Data Retention",
        "risk": "Conversation history stored in AgentCore Memory (STM + LTM) could contain sensitive operational details or PII",
        "impact": "Data retention beyond operator expectations, potential compliance issues",
        "likelihood": "Medium — Memory is enabled by default. No PII filtering is applied before storage",
    },
    {
        "title": "OpenSearch Serverless Wildcard Access",
        "risk": "The Knowledge Base IAM role has aoss:APIAccessAll on all collections (collection/*)",
        "impact": "If the role is compromised, it could access any OpenSearch Serverless collection in the account",
        "likelihood": "Low — Role is only assumable by the Bedrock service principal",
    },
]

for threat in threats:
    doc.add_heading(threat["title"], level=2)
    items = [
        ("Risk:", threat["risk"]),
        ("Impact:", threat["impact"]),
        ("Likelihood:", threat["likelihood"]),
    ]
    for bold_text, normal_text in items:
        p = doc.add_paragraph(style='List Bullet')
        run_bold = p.add_run(bold_text + " ")
        run_bold.bold = True
        run_normal = p.add_run(normal_text)

# --- Section 3 ---
doc.add_heading('3. What can we do about it?', level=1)

doc.add_heading('Access Control', level=2)
access_controls = [
    "All API endpoints require Cognito JWT authorizer (data endpoints, chat, voice-token)",
    "Cognito User Pool: admin-only registration (AllowAdminCreateUserOnly: true), self-signup disabled",
    "Password policy: 12+ characters, uppercase, lowercase, numbers, symbols required",
    "MFA enforced via SOFTWARE_TOKEN_MFA (MfaConfiguration: ON)",
    "Cognito AdvancedSecurityMode set to ENFORCED for adaptive authentication and compromised credential detection",
    "API Lambda IAM role scoped to five specific DynamoDB table ARNs with aws:RequestedRegion and aws:SourceAccount conditions",
    "Chat Lambda IAM role scoped to AgentCore invocation and SQS DLQ",
    "CloudFront Origin Access Control — S3 dashboard bucket only accessible through CDN",
    "S3 public access blocked on all buckets (dashboard, access logs, KB data)",
    "API Gateway request validation enabled",
]
for item in access_controls:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Data Protection', level=2)
data_protection = [
    "All 6 DynamoDB tables encrypted with KMS Customer Managed Key (automatic annual rotation)",
    "Lambda environment variables encrypted with separate KMS CMK (key isolation)",
    "CloudFront enforces HTTPS with TLS 1.2 minimum (TLSv1.2_2021)",
    "Security response headers: HSTS (max-age=31536000), X-Frame-Options: DENY, X-Content-Type-Options: nosniff, Referrer-Policy: strict-origin-when-cross-origin",
    "S3 bucket policies enforce SSL-only access (aws:SecureTransport deny condition)",
    "SQS dead letter queue policies enforce SSL-only access",
    "Chat-history DynamoDB table has TTL for automatic conversation expiry",
    "DynamoDB Point-in-Time Recovery enabled on all tables",
    "S3 versioning enabled on dashboard and KB data buckets",
    "Agent system prompt defines strict data boundaries — tools only return DynamoDB data",
    "Lambda error handlers return generic messages, no stack traces exposed",
]
for item in data_protection:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Infrastructure Security', level=2)
infra_security = [
    "AWS WAF on CloudFront with AWS Managed Rules: AWSManagedRulesCommonRuleSet and AWSManagedRulesKnownBadInputsRuleSet",
    "Lambda reserved concurrency (10 per function) prevents noisy-neighbor exhaustion",
    "SQS dead letter queues capture failed Lambda invocations for analysis",
    "CloudFront access logging to dedicated S3 bucket with 90-day lifecycle",
    "CloudWatch Logs configured for all Lambda functions",
    "CloudTrail enabled for API call auditing",
    "All infrastructure deployed via CloudFormation templates",
]
for item in infra_security:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Code Security', level=2)

doc.add_paragraph('ASH v3.2.3 scan completed (March 12, 2026) with 10 scanners:')

code_findings = [
    "Bandit (Python SAST): PASSED — 0 actionable findings (4 low: random.uniform usage in data loader)",
    "cdk-nag (CloudFormation): PASSED — 25 original findings resolved. Cognito MFA and AdvancedSecurityMode remediated. Remaining are false positives (SQS2, KMS5, CFR4) or by design (APIG4 OPTIONS, IAM5 AgentCore)",
    "Checkov (IaC security): PASSED — 8 original findings resolved. KB S3 logging and versioning remediated. Remaining are by design (Lambda VPC, SQS encryption, IAM wildcards)",
    "detect-secrets: PASSED — 1 finding confirmed as false positive (GenerateSecret: false in Cognito client config)",
    "Semgrep (code patterns): PASSED — 6 original findings resolved. SRI integrity added to 3d-twin.html (2), unsafe-formatstring fixed in api.js (1). Remaining innerHTML in shared.js is by design with escapeHtml sanitizer",
    "npm-audit: PASSED — no dependency vulnerabilities",
]
for item in code_findings:
    doc.add_paragraph(item, style='List Bullet')

additional_code = [
    "Agent implements 3-retry mechanism with error classification for transient Bedrock failures",
    "Input validation on API Lambda (path length check, method validation)",
    "No hardcoded credentials in application code — table names and ARNs via CloudFormation environment variables",
]
for item in additional_code:
    doc.add_paragraph(item, style='List Bullet')

# --- Section 4 ---
doc.add_heading('4. Did we do a good enough job (for now)?', level=1)

doc.add_heading('Review & Validation', level=2)
review_items = [
    "All infrastructure code scanned with ASH v3.2.3 (Bandit, cdk-nag, Checkov, detect-secrets, Semgrep, npm-audit)",
    "CloudFormation templates analyzed by cdk-nag with AwsSolutionsChecks enabled",
    "Python agent code analyzed by Bandit and Semgrep",
    "Architecture reviewed against STRIDE threat categories",
    "Well-Architected Framework assessment completed across all 6 pillars",
    "The solution is an AWS Guidance intended as a reference architecture and demo, not a production-ready deployment as-is",
    "DSR (Deliverable Security Review) has been completed alongside this threat model",
]
for item in review_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Residual Risk', level=2)
residual_risks = [
    "Bedrock Guardrails not configured: No content filters, prompt attack detection, or PII filters are enabled. Customers MUST configure Bedrock Guardrails before production use",
    "CORS wildcard: Access-Control-Allow-Origin: * is set on all API responses. Customers MUST restrict to their CloudFront domain before production",
    "API Gateway throttling not configured: No rate/burst limits beyond WAF managed rules. Customers should add throttling settings for production",
    "Demo credentials in deployment script: deploy-all.sh contains hardcoded demo user email and password. Customers should remove and use Secrets Manager",
    "IAM wildcards on AgentCore: Resource: * for AgentCore invocation permissions. Customers should scope to specific agent ARN after deployment",
    "OpenSearch Serverless wildcard: aoss:APIAccessAll on collection/*. Customers should scope to specific collection ARN",
    "CloudWatch Alarms not configured: No automated alerting on Lambda errors, API 4xx/5xx, or DynamoDB throttling. Customers should add alarms for production",
    "Conversation memory PII: No PII filtering before storing conversation history in AgentCore Memory. Customers handling PII should implement filtering",
    "ASH findings after remediation: 40 original findings reduced by fixes. Remaining are false positives or by-design: cdk-nag SQS2 (SqsManagedSseEnabled set), KMS5 (EnableKeyRotation set), CFR4 (TLSv1.2_2021 set with default cert), APIG4 on OPTIONS (CORS preflight), IAM5 on AgentCore (dynamic ARNs), detect-secrets (GenerateSecret: false), Semgrep innerHTML (escapeHtml sanitizer), Checkov CKV_AWS_117 (serverless, no VPC needed)",
]
for item in residual_risks:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'Given these constraints and the solution\'s purpose as an AWS Guidance (reference architecture), '
    'we have done enough to protect customers who decide to evaluate and adapt this solution. '
    'Customers must address the residual risks above before any production deployment.'
)

# --- Notes ---
doc.add_heading('Notes', level=1)
notes = [
    "This threat model was generated on 2026-03-15 and updated after security remediation based on ASH v3.2.3 scan results",
    "The threat model should be updated when the architecture changes or new features are added",
    "Customers should conduct their own security assessment and threat modeling before deploying to production",
    "Additional security measures (Bedrock Guardrails, enforced MFA, CORS restriction, API throttling, WAF rate limiting, VPC endpoints) are recommended for production use",
]
for item in notes:
    doc.add_paragraph(item, style='List Bullet')

# Save
doc.save('temp/Threat Model.docx')
print("Done: temp/Threat Model.docx")
